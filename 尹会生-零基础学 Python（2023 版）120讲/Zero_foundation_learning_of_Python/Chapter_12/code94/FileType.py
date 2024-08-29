# NOTE: Word类型：python-docx
from docx import Document
from docx.shared import Inches

doc = Document()

doc.add_heading("标题", 0)
# 加粗
p = doc.add_paragraph('A plain paragraph having some')
p.add_run('bold').bold = True
p.add_run(' and some')
p.add_run('italic.').italic = True

# 标题
doc.add_heading('Heading leval 1', level=1)

# 内容
doc.add_paragraph('Intense quote', style='Intense Quote')

doc.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
doc.add_paragraph(
    'first item in ordered list', style='List Number'
)

# 图片
# doc.add_picture('monty-truth.png', width=Inches(1.25))

# 表格
records = {
    (3, '101', 'Spam'), (7, '422', 'Eggs'), (4, '631', 'Spam, spam, eggs, and spam')
}
table = doc.add_table(rows=1, cols=3)
# 表格首行
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'

for Qty, Id, Desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(Qty)
    row_cells[1].text = Id
    row_cells[2].text = Desc

doc.add_page_break()
# 保存文件名
doc.save('simple.docx')

# NOTE: excel类型:openpyxl
from openpyxl import Workbook

wb = Workbook()

ws = wb.active
ws['A1'] = 42
ws.append([1, 2, 3])
import datetime

ws['A2'] = datetime.datetime.now()
wb.save('simple.xlsx')
